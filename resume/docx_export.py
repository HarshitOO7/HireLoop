"""
Markdown → ATS-safe .docx via python-docx.

ATS rules enforced:
  - Single column, NO tables, NO text boxes, NO images
  - Standard paragraph styles only (Normal, List Bullet)
  - Section headings as bold paragraph + bottom border (not heading styles —
    some ATS parsers choke on non-standard heading fonts)
  - Standard font: Calibri (default in python-docx template)
  - Readable section header text: WORK EXPERIENCE, SKILLS, etc.

Handles the Markdown format produced by ai.tailor_resume():
  # Name            → large bold centered name
  contact line      → centered normal text (contains | separators)
  ## SECTION        → bold ALL-CAPS + bottom border rule
  **Role** | Date   → bold + normal runs in one paragraph
  *Company | Loc*   → italic paragraph
  - bullet          → List Bullet style
  **Key:** text     → bold key + normal value inline
  blank line        → skip (no extra spacing)

Usage:
  from resume.docx_export import render_docx
  path = render_docx(markdown_text, "resume/output/resume.docx")
"""

import re
import logging
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

logger = logging.getLogger(__name__)

# ── Sizes ─────────────────────────────────────────────────────────────────────
_NAME_PT    = Pt(20)
_SECTION_PT = Pt(11)
_BODY_PT    = Pt(10.5)


# ── XML helpers ───────────────────────────────────────────────────────────────

def _bottom_border(paragraph):
    """Add a thin bottom border under a paragraph (section divider)."""
    pPr    = paragraph._p.get_or_add_pPr()
    pBdr   = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _spacing(paragraph, before: int = 0, after: int = 2):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after  = Pt(after)


# ── Contact line helpers ──────────────────────────────────────────────────────

_URL_DISPLAY: dict[str, str] = {
    "linkedin.com": "LinkedIn",
    "github.com":   "GitHub",
    "behance.net":  "Behance",
    "dribbble.com": "Dribbble",
    "kaggle.com":   "Kaggle",
}

_RE_URL = re.compile(r"(https?://)?(([\w\-]+\.)+[\w]{2,})(/[\w\-./%~:@!$&'()*+,;=?#]*)?")


def _url_display_text(url: str) -> str:
    """Return a short display label for a URL (e.g. 'LinkedIn', 'GitHub', or bare domain)."""
    # Strip scheme for matching
    clean = re.sub(r"^https?://", "", url).lstrip("www.")
    for domain, label in _URL_DISPLAY.items():
        if clean.startswith(domain):
            return label
    # Fallback: just the domain part
    return clean.split("/")[0]


def _add_hyperlink(paragraph, url: str, display_text: str, size: Pt = _BODY_PT) -> None:
    """Add a clickable hyperlink run to a paragraph. No blue/underline styling — clean look."""
    if not url.startswith("http"):
        url = "https://" + url
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run_elem = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    # Suppress default hyperlink style (blue underline)
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    # Override color to auto (black) and remove underline
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "auto")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "none")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size.pt * 2)))
    rPr.append(color)
    rPr.append(u)
    rPr.append(sz)
    run_elem.append(rPr)
    t = OxmlElement("w:t")
    t.text = display_text
    run_elem.append(t)
    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)


def _add_contact_line(paragraph, line: str, size: Pt = _BODY_PT) -> None:
    """
    Render a contact line (phone | email | linkedin.com/... | github.com/...).
    URL tokens become real DOCX hyperlinks with clean display text.
    Plain tokens (phone, email) are plain runs.
    """
    tokens = [t.strip() for t in line.split("|")]
    for i, token in enumerate(tokens):
        if i > 0:
            run = paragraph.add_run(" | ")
            run.font.size = size
        if _RE_URL.fullmatch(token.strip()):
            _add_hyperlink(paragraph, token.strip(), _url_display_text(token.strip()), size)
        else:
            run = paragraph.add_run(token)
            run.font.size = size


# ── Inline markup parser ──────────────────────────────────────────────────────

def _add_inline(paragraph, text: str, default_size: Pt = _BODY_PT):
    """
    Parse **bold** and *italic* markup and add runs to `paragraph`.
    Plain text → normal run.
    """
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold      = True
            run.font.size = default_size
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic    = True
            run.font.size = default_size
        else:
            run = paragraph.add_run(part)
            run.font.size = default_size


# ── Main export ───────────────────────────────────────────────────────────────

def render_docx(markdown_text: str, output_path: str | Path) -> Path:
    """
    Convert a Markdown resume string to an ATS-safe .docx file.
    Returns the output Path.
    """
    doc = Document()

    # ── Margins: 0.75" sides, 0.5" top/bottom ─────────────────────────────
    for sec in doc.sections:
        sec.top_margin    = Pt(36)
        sec.bottom_margin = Pt(36)
        sec.left_margin   = Pt(54)
        sec.right_margin  = Pt(54)

    # Remove the blank paragraph that Document() always creates
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    lines = markdown_text.splitlines()
    first_line_done = False   # track whether we've seen the name yet

    for raw_line in lines:
        line = raw_line.rstrip()

        # ── Skip blank lines ─────────────────────────────────────────────────
        if not line.strip():
            continue

        # ── Strip HTML comments (e.g. section-order hint injected by generator) ──
        if line.strip().startswith("<!--"):
            continue

        # ── # Name — large bold centered ─────────────────────────────────────
        if line.startswith("# "):
            name = line[2:].strip()
            p    = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _spacing(p, before=0, after=2)
            run = p.add_run(name)
            run.bold      = True
            run.font.size = _NAME_PT
            first_line_done = True
            continue

        # ── ## Section heading — bold ALL-CAPS + bottom border ───────────────
        if line.startswith("## "):
            title = line[3:].strip().upper()
            p     = doc.add_paragraph()
            _spacing(p, before=6, after=1)
            run = p.add_run(title)
            run.bold      = True
            run.font.size = _SECTION_PT
            _bottom_border(p)
            continue

        # ── Bullet point: - text or * text ───────────────────────────────────
        if re.match(r"^[-*] ", line):
            content = line[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            _spacing(p, before=0, after=1)
            p.paragraph_format.left_indent = Pt(18)
            _add_inline(p, content)
            continue

        # ── *Italic line* — company / subtitle (whole line italic) ───────────
        if (
            line.startswith("*") and line.endswith("*")
            and not line.startswith("**")
            and line.count("*") == 2
        ):
            p = doc.add_paragraph()
            _spacing(p, before=0, after=0)
            run = p.add_run(line[1:-1])
            run.italic    = True
            run.font.size = _BODY_PT
            continue

        # ── Contact line (contains | separators, no leading **) — centered ───
        if "|" in line and not line.startswith("**") and not first_line_done is False:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _spacing(p, before=0, after=4)
            _add_contact_line(p, line, size=_BODY_PT)
            continue

        # ── Everything else — inline parse (role lines, **Key:** val, etc.) ──
        p = doc.add_paragraph()
        _spacing(p, before=1, after=1)
        _add_inline(p, line)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("[docx_export] saved → %s  (%d lines processed)", output_path, len(lines))
    return output_path
